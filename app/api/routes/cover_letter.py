from io import BytesIO
from uuid import UUID

import structlog
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.config import current_active_user
from app.db.dependencies import get_db
from app.db.models import ApplicationModel, User
from app.llm.cover_letter import CoverLetterGenerator
from app.llm.exceptions import (
    LLMAllKeysExhaustedError,
    LLMRateLimitError,
    LLMValidationError,
)
from app.schemas.cover_letter import CoverLetterResponse
from app.schemas.tailored_cv import TailoredCV

router = APIRouter()
logger = structlog.get_logger()


@router.post("/applications/{application_id}/cover-letter", response_model=CoverLetterResponse)
async def generate_cover_letter(
    application_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(current_active_user),
):
    result = await db.execute(
        select(ApplicationModel).where(ApplicationModel.id == application_id)
    )
    application = result.scalar_one_or_none()
    if application is None:
        raise HTTPException(status_code=404, detail="Application not found")
    if application.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden")

    tailored_cv = TailoredCV.model_validate(application.tailored_cv_data)
    generator = CoverLetterGenerator()

    try:
        cover_letter_text = await generator.generate(tailored_cv, application.job_description)
    except LLMAllKeysExhaustedError:
        raise HTTPException(
            status_code=503,
            detail="Daily usage limit reached. The service resets at midnight.",
        )
    except LLMRateLimitError as e:
        raise HTTPException(
            status_code=429,
            detail=f"Service is busy. Please try again in {e.retry_after_seconds} seconds.",
        )
    except LLMValidationError as e:
        raise HTTPException(status_code=422, detail=str(e))

    application.cover_letter = cover_letter_text
    await db.commit()

    await logger.ainfo("cover_letter_saved", application_id=str(application_id))
    return CoverLetterResponse(cover_letter=cover_letter_text)


@router.get("/applications/{application_id}/cover-letter/download")
async def download_cover_letter(
    application_id: UUID,
    format: str = Query(default="txt"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(current_active_user),
):
    if format not in ("txt", "docx", "pdf"):
        raise HTTPException(status_code=422, detail="format must be one of: txt, docx, pdf")

    result = await db.execute(
        select(ApplicationModel).where(ApplicationModel.id == application_id)
    )
    application = result.scalar_one_or_none()
    if application is None:
        raise HTTPException(status_code=404, detail="Application not found")
    if application.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden")
    if not application.cover_letter:
        raise HTTPException(status_code=404, detail="Cover letter not yet generated")

    letter_text = application.cover_letter
    full_name = (application.tailored_cv_data or {}).get("full_name", "cover_letter")
    safe_name = full_name.replace(" ", "_").lower()

    if format == "txt":
        content = letter_text.encode("utf-8")
        filename = f"{safe_name}_cover_letter.txt"
        return StreamingResponse(
            BytesIO(content),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if format == "docx":
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        title_para = doc.add_paragraph()
        run = title_para.add_run(f"{full_name} — Cover Letter")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()  # blank spacer line
        paragraphs = [p.strip() for p in letter_text.split("\n\n") if p.strip()]
        for para in paragraphs:
            doc.add_paragraph(para)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{safe_name}_cover_letter.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # pdf
    from weasyprint import HTML

    paragraphs = [p.strip() for p in letter_text.split("\n\n") if p.strip()]
    para_html = "".join(f"<p>{p}</p>" for p in paragraphs)
    html_content = f"""<!DOCTYPE html>
<html>
<head>
  <style>
    body {{ font-family: Georgia, serif; margin: 2cm; line-height: 1.6; color: #222; }}
    h1 {{ font-size: 14pt; margin-bottom: 1.5em; }}
    p {{ margin-bottom: 1em; font-size: 11pt; text-align: justify; }}
  </style>
</head>
<body>
  <h1>{full_name} — Cover Letter</h1>
  {para_html}
</body>
</html>"""
    pdf_bytes = HTML(string=html_content).write_pdf()
    filename = f"{safe_name}_cover_letter.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
