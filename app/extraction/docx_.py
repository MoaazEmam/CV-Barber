from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn


class DocxExtractor:
    def extract_text(self, path: str) -> str:
        doc = Document(path)
        lines: list[str] = []

        # Headers/footers first — contact details often live here in templates.
        for section in doc.sections:
            for container in (section.header, section.footer):
                lines.extend(self._container_lines(container))

        # Main body: paragraphs + tables (the original behaviour).
        lines.extend(self._container_lines(doc))

        # Text boxes / shapes (w:txbxContent). Modern text boxes appear twice in
        # the XML (DrawingML + a VML fallback), so dedupe identical contents.
        seen_box: set[str] = set()
        for txbx in doc.element.body.iter(qn("w:txbxContent")):
            text = self._element_text(txbx)
            if text and text not in seen_box:
                seen_box.add(text)
                lines.append(text)

        text = "\n".join(line for line in lines if line)

        # Real hyperlink targets (LinkedIn/GitHub/website) — paragraph.text drops
        # the URL, keeping only the visible label.
        links = self._collect_links(doc)
        if links:
            text = f"{text}\nLinks: {' '.join(links)}"

        return text

    def _container_lines(self, container) -> list[str]:
        out: list[str] = []
        for paragraph in container.paragraphs:
            t = paragraph.text.strip()
            if t:
                out.append(t)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        out.append(t)
        return out

    def _element_text(self, element) -> str:
        parts = [node.text for node in element.iter(qn("w:t")) if node.text]
        return " ".join(p.strip() for p in parts if p.strip())

    def _collect_links(self, doc) -> list[str]:
        seen: list[str] = []
        parts = [doc.part]
        for section in doc.sections:
            parts.extend([section.header.part, section.footer.part])
        for part in parts:
            try:
                rels = part.rels.values()
            except Exception:
                continue
            for rel in rels:
                if rel.reltype == RT.HYPERLINK and rel.is_external:
                    uri = rel.target_ref
                    if uri and uri not in seen:
                        seen.append(uri)
        return seen
