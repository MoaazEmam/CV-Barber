from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from app.generation.base_generator import BaseGenerator
from app.generation.styles import (
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
    FONT_NAME, FONT_SIZE_NAME, FONT_SIZE_CONTACT, FONT_SIZE_SECTION,
    FONT_SIZE_BODY, FONT_SIZE_ENTRY_TITLE,
    SPACE_BEFORE_SECTION, SPACE_AFTER_SECTION,
    SPACE_AFTER_ENTRY, SPACE_AFTER_BULLET,
    ALIGN_CENTER,
)
from app.schemas.tailored_cv import TailoredCV, ScoredExperienceEntry, ScoredProjectEntry


class DocxGenerator(BaseGenerator):
    def generate(self, cv: TailoredCV) -> bytes:
        doc = Document()
        self._set_margins(doc)
        self._add_header(doc, cv)
        self._add_education(doc, cv)
        self._add_skills(doc, cv)
        self._add_experience_and_projects(doc, cv)
        return self._to_bytes(doc)

    def content_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def file_extension(self) -> str:
        return "docx"

    def _set_margins(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    def _add_header(self, doc: Document, cv: TailoredCV) -> None:
        # name
        name_para = doc.add_paragraph()
        name_para.alignment = ALIGN_CENTER
        name_para.paragraph_format.space_after = Pt(2)
        run = name_para.add_run(cv.full_name)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NAME

        # contact line
        contact_para = doc.add_paragraph()
        contact_para.alignment = ALIGN_CENTER
        contact_para.paragraph_format.space_after = Pt(1)

        parts = []
        if cv.phone:
            parts.append(("text", cv.phone))
        if cv.email:
            if parts:
                parts.append(("text", " | "))
            parts.append(("link", cv.email, f"mailto:{cv.email}"))
        if cv.linkedin:
            if parts:
                parts.append(("text", " | "))
            url = cv.linkedin if cv.linkedin.startswith("http") else f"http://{cv.linkedin}"
            parts.append(("link", cv.linkedin, url))

        for part in parts:
            if part[0] == "text":
                run = contact_para.add_run(part[1])
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_CONTACT
            else:
                self._add_hyperlink(contact_para, part[1], part[2])

        # github
        if cv.github:
            github_para = doc.add_paragraph()
            github_para.alignment = ALIGN_CENTER
            github_para.paragraph_format.space_after = Pt(4)
            url = cv.github if cv.github.startswith("http") else f"https://{cv.github}"
            self._add_hyperlink(github_para, cv.github, url)

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """
        Inserts a real clickable hyperlink into a paragraph via OOXML.
        python-docx doesn't have a native hyperlink API so we use XML directly.
        """
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run_elem = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        # font name
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)
        rPr.append(rFonts)

        # font size (half-points)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(FONT_SIZE_CONTACT.pt * 2)))
        rPr.append(sz)

        run_elem.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run_elem.append(t)
        hyperlink.append(run_elem)
        paragraph._p.append(hyperlink)

    def _add_section_heading(self, doc: Document, title: str) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = SPACE_BEFORE_SECTION
        para.paragraph_format.space_after = SPACE_AFTER_SECTION
        run = para.add_run(title)
        run.bold = False
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_SECTION

        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_education(self, doc: Document, cv: TailoredCV) -> None:
        self._add_section_heading(doc, "EDUCATION")

        for entry in cv.education:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(1)
            self._set_right_tab(para)

            parts = [entry.institution]
            if entry.faculty:
                parts.append(entry.faculty)
            if entry.gpa:
                parts.append(f"GPA: {entry.gpa}")

            run = para.add_run(" ||".join(parts))
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

            date_str = self._format_date_range(entry.date_range)
            tab_run = para.add_run(f"\t{date_str}")
            tab_run.bold = False
            tab_run.font.name = FONT_NAME
            tab_run.font.size = FONT_SIZE_BODY

            major_para = doc.add_paragraph()
            major_para.paragraph_format.space_after = Pt(2)
            run = major_para.add_run(f"Major: {entry.field}")
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

    def _add_skills(self, doc: Document, cv: TailoredCV) -> None:
        self._add_section_heading(doc, "SKILLS")

        for category in cv.skills:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = SPACE_AFTER_BULLET

            run = para.add_run(f"{category.category}: ")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

            run = para.add_run(", ".join(category.skills))
            run.bold = False
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

        # certifications as a bullet with nested items
        if cv.certifications:
            cert_para = doc.add_paragraph(style="List Bullet")
            cert_para.paragraph_format.space_after = SPACE_AFTER_BULLET
            run = cert_para.add_run("Certifications & Courses:")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

            for cert in cv.certifications:
                nested = doc.add_paragraph(style="List Bullet 2")
                nested.paragraph_format.space_after = SPACE_AFTER_BULLET
                run = nested.add_run(cert)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_BODY

    def _add_experience_and_projects(self, doc: Document, cv: TailoredCV) -> None:
        if cv.experience:
            self._add_section_heading(doc, "EXPERIENCE")
            for entry in cv.experience:
                self._add_experience_entry(doc, entry)

        if cv.projects:
            self._add_section_heading(doc, "PROJECTS")
            for entry in cv.projects:
                self._add_project_entry(doc, entry)

    def _add_experience_entry(self, doc: Document, entry: ScoredExperienceEntry) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        self._set_right_tab(para)

        title_run = para.add_run(f"{entry.title}:")
        title_run.bold = True
        title_run.font.name = FONT_NAME
        title_run.font.size = FONT_SIZE_ENTRY_TITLE

        company_parts = []
        if entry.company:
            company_parts.append(entry.company)
        if entry.location:
            company_parts.append(f"({entry.location})")

        company_run = para.add_run(f" {' '.join(company_parts)}")
        company_run.italic = True
        company_run.font.name = FONT_NAME
        company_run.font.size = FONT_SIZE_ENTRY_TITLE

        date_str = self._format_date_range(entry.date_range)
        date_run = para.add_run(f"\t{date_str}")
        date_run.font.name = FONT_NAME
        date_run.font.size = FONT_SIZE_ENTRY_TITLE

        for bullet in entry.bullets:
            self._add_bullet(doc, bullet)

        doc.paragraphs[-1].paragraph_format.space_after = SPACE_AFTER_ENTRY

    def _add_project_entry(self, doc: Document, entry: ScoredProjectEntry) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        self._set_right_tab(para)

        name_run = para.add_run(f"{entry.name}:")
        name_run.bold = True
        name_run.font.name = FONT_NAME
        name_run.font.size = FONT_SIZE_ENTRY_TITLE

        if entry.description:
            desc_run = para.add_run(f" {entry.description}")
            desc_run.italic = True
            desc_run.font.name = FONT_NAME
            desc_run.font.size = FONT_SIZE_ENTRY_TITLE

        if entry.date_range and entry.date_range.start:
            date_run = para.add_run(f"\t{entry.date_range.start}")
            date_run.font.name = FONT_NAME
            date_run.font.size = FONT_SIZE_ENTRY_TITLE

        for bullet in entry.bullets:
            self._add_bullet(doc, bullet)

        doc.paragraphs[-1].paragraph_format.space_after = SPACE_AFTER_ENTRY

    def _add_bullet(self, doc: Document, text: str) -> None:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = SPACE_AFTER_BULLET
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY

    def _set_right_tab(self, para) -> None:
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        # page width 8.5in - 0.25in*2 margins = 8in = 11520 twips
        tab.set(qn("w:pos"), "11520")
        tabs.append(tab)
        pPr.append(tabs)

    def _format_date_range(self, date_range) -> str:
        if not date_range:
            return ""
        if date_range.end:
            return f"{date_range.start} – {date_range.end}"
        return date_range.start if date_range.start else ""

    def _to_bytes(self, doc: Document) -> bytes:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()