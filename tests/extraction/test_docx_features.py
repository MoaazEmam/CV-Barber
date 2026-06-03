"""Tests for DOCX header/footer, text box, and hyperlink-target extraction."""

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.extraction.docx_ import DocxExtractor


def _add_hyperlink(paragraph, url: str, label: str) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = label
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_textbox(doc, text: str) -> None:
    # Minimal w:txbxContent subtree (all in the 'w' namespace) appended to the
    # body — enough to exercise the extractor's text-box walk.
    txbx = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    txbx.append(p)
    doc.element.body.append(txbx)


def test_header_textbox_and_hyperlink(tmp_path):
    doc = Document()

    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.add_paragraph("Jane Doe — jane@example.com")

    para = doc.add_paragraph("Find me on ")
    _add_hyperlink(para, "https://github.com/janedoe", "GitHub")

    _add_textbox(doc, "Skills: Python, Rust, Go")

    path = tmp_path / "cv.docx"
    doc.save(str(path))

    text = DocxExtractor().extract_text(str(path))
    assert "Jane Doe" in text                       # header
    assert "jane@example.com" in text               # header
    assert "Skills: Python, Rust, Go" in text       # text box
    assert "https://github.com/janedoe" in text     # hyperlink target
    assert "Links:" in text


def test_plain_docx_still_extracts(tmp_path):
    doc = Document()
    doc.add_paragraph("Senior Engineer at Acme")
    doc.add_paragraph("Built distributed systems")
    path = tmp_path / "plain.docx"
    doc.save(str(path))

    text = DocxExtractor().extract_text(str(path))
    assert "Senior Engineer at Acme" in text
    assert "Built distributed systems" in text
