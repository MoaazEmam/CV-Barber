import io

from docx import Document

from app.pipeline.docx.render import render
from app.schemas.tailored_cv import TailoredCV


def _doc_bytes(paragraphs: list[str], bullet_idx=()) -> bytes:
    # Real CVs format bullets as Word list items; the renderer only rewrites those
    # (not structural sub-lines), so the test fixtures must mark bullets as such.
    doc = Document()
    for i, text in enumerate(paragraphs):
        if i in bullet_idx:
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _texts(docx_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def _exp(title, company, bullets, score=8):
    return {
        "title": title,
        "company": company,
        "date_range": {"start": "2020"},
        "bullets": bullets,
        "relevance_score": score,
        "relevance_reason": "r",
    }


def _tailored(experience) -> TailoredCV:
    return TailoredCV.model_validate(
        {
            "full_name": "John Doe",
            "tailored_summary": "New summary",
            "job_title": "Engineer",
            "company_name": "Acme",
            "experience": experience,
            "projects": [],
        }
    )


# Paragraph layout used across tests:
#   0 John Doe | 1 Summary stub | 2 Experience
#   3 Backend Engineer at A | 4 Built X
#   5 Frontend Engineer at B | 6 Made Y
def _layout_bytes():
    return _doc_bytes([
        "John Doe",
        "Old summary",
        "Experience",
        "Backend Engineer at A",
        "Built X",
        "Frontend Engineer at B",
        "Made Y",
    ], bullet_idx={4, 6})


def _artifact():
    return {
        "format": "docx",
        "summary_paragraphs": [1],
        "experience": [
            {"paragraphs": [3, 4], "heading": 3},
            {"paragraphs": [5, 6], "heading": 5},
        ],
        "projects": [],
        "education": [],
    }


def test_drops_unselected_entry_and_rewrites_bullet():
    tailored = _tailored([_exp("Frontend Engineer", "B", ["Rewrote Y"])])
    out = _texts(render(_layout_bytes(), _artifact(), tailored))

    assert "Frontend Engineer at B" in out
    assert "Rewrote Y" in out
    # Dropped entry and its bullet are gone.
    assert "Backend Engineer at A" not in out
    assert "Built X" not in out
    assert "Made Y" not in out


def test_summary_replaced():
    tailored = _tailored([_exp("Backend Engineer", "A", ["Built X"])])
    out = _texts(render(_layout_bytes(), _artifact(), tailored))
    assert "New summary" in out
    assert "Old summary" not in out


def test_preserves_structural_subline_and_only_rewrites_bullets():
    # Entry = heading + a non-bullet sub-line (company/location) + one real bullet.
    layout = _doc_bytes([
        "John Doe", "Old summary", "Experience",
        "Backend Engineer at A",  # 3 heading
        "Cairo, Egypt",           # 4 structural sub-line (NOT a list item)
        "Built X",                # 5 bullet
    ], bullet_idx={5})
    artifact = {
        "format": "docx", "summary_paragraphs": [1],
        "experience": [{"paragraphs": [3, 4, 5], "heading": 3}],
        "projects": [], "education": [],
    }
    tailored = _tailored([_exp("Backend Engineer", "A", ["Rebuilt X"])])
    out = _texts(render(layout, artifact, tailored))
    assert "Cairo, Egypt" in out   # sub-line preserved, not overwritten with bullet text
    assert "Rebuilt X" in out      # the real bullet was rewritten
    assert "Built X" not in out


def test_matches_correct_block_despite_techstack_suffix():
    # Headings carry "Name - TechStack - Date". The correct (longer) heading must
    # win over a shorter heading that merely shares a word like "Management System"
    # — the bug that put Education-Ops bullets under the Fashion-House heading.
    layout = _doc_bytes([
        "Jane", "Sum", "PROJECTS",
        "Education Operations Management System - Microsoft Excel (advanced automation, dashboards)",  # 3
        "Edu original",  # 4 bullet
        "Fashion House Management System - C#, .NET, SQL Server 2024",  # 5
        "Fashion original",  # 6 bullet
    ], bullet_idx={4, 6})
    artifact = {
        "format": "docx", "summary_paragraphs": [1], "experience": [], "education": [],
        "projects": [
            {"paragraphs": [3, 4], "heading": 3},
            {"paragraphs": [5, 6], "heading": 5},
        ],
    }
    tailored = TailoredCV.model_validate({
        "full_name": "Jane", "job_title": "E", "company_name": "C", "experience": [],
        "projects": [
            {"name": "Education Operations Management System", "bullets": ["Edu bullet"],
             "relevance_score": 8, "relevance_reason": "r"},
            {"name": "Fashion House Management System", "bullets": ["Fashion bullet"],
             "relevance_score": 7, "relevance_reason": "r"},
        ],
    })
    out = _texts(render(layout, artifact, tailored))
    assert "Education Operations Management System" in "\n".join(out)
    assert "Fashion House Management System" in "\n".join(out)
    # Each bullet sits under its OWN heading (no cross-wiring).
    i_edu_h = next(i for i, t in enumerate(out) if t.startswith("Education Operations"))
    i_edu_b = out.index("Edu bullet")
    i_fash_h = next(i for i, t in enumerate(out) if t.startswith("Fashion House"))
    i_fash_b = out.index("Fashion bullet")
    assert i_edu_h < i_edu_b < i_fash_h < i_fash_b


def test_reorders_kept_entries():
    # Keep both, but in reversed order: Frontend first, Backend second.
    tailored = _tailored([
        _exp("Frontend Engineer", "B", ["Made Y"]),
        _exp("Backend Engineer", "A", ["Built X"]),
    ])
    out = _texts(render(_layout_bytes(), _artifact(), tailored))
    assert out.index("Frontend Engineer at B") < out.index("Backend Engineer at A")
