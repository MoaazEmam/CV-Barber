import io

import fitz
from docx import Document

from app.pipeline.docx.parse import parse_spans as docx_spans
from app.pipeline.pdf.parse import parse_spans as pdf_spans, render_pages_png


def _make_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Jane Smith", fontsize=18)
    page.insert_text((72, 110), "Software Engineer", fontsize=11)
    return doc.tobytes()


def test_pdf_spans_roundtrip():
    spans = pdf_spans(_make_pdf_bytes())
    texts = [s.text for s in spans]
    assert "Jane Smith" in texts
    assert "Software Engineer" in texts
    # ids are unique and the prompt dict carries the text.
    assert len({s.id for s in spans}) == len(spans)
    assert spans[0].to_prompt_dict()["text"] == spans[0].text


def test_render_pages_png_returns_png_bytes():
    pngs = render_pages_png(_make_pdf_bytes())
    assert len(pngs) == 1
    assert pngs[0][:8] == b"\x89PNG\r\n\x1a\n"


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("John Doe")
    p = doc.add_paragraph("")
    run = p.add_run("Senior Engineer")
    run.bold = True
    doc.add_paragraph("Did things")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_docx_spans_encode_paragraph_index_and_bold():
    spans = docx_spans(_make_docx_bytes())
    by_text = {s.text: s for s in spans}
    assert "John Doe" in by_text
    # span id encodes the paragraph index as "p{n}".
    assert by_text["John Doe"].id.startswith("p")
    assert by_text["Senior Engineer"].bold is True
