# tests/generation/test_docx_generator.py
import pytest
from docx import Document
from io import BytesIO
from app.generation.docx_generator import DocxGenerator
from app.schemas.cv_blocks import DateRange
from app.schemas.tailored_cv import (
    TailoredCV, ScoredExperienceEntry, ScoredProjectEntry,
    EducationEntry, SkillCategory
)


@pytest.fixture
def sample_tailored_cv() -> TailoredCV:
    return TailoredCV(
        full_name="Moaaz Emam Ahmed",
        email="moaaz@example.com",
        phone="+20 1156637443",
        github="https://github.com/MoaazEmam",
        linkedin="www.linkedin.com/in/moaaz-emam",
        job_title="Backend Engineering Intern",
        company_name="Acme Corp",
        tailored_summary="Experienced backend developer.",
        education=[EducationEntry(
            institution="Cairo University",
            faculty="Faculty of Engineering",
            field="Communication and Computer Engineering",
            date_range=DateRange(start="2022", end="2027"),
            gpa="3.58",
        )],
        skills=[SkillCategory(
            category="Programming Languages",
            skills=["Python", "TypeScript"]
        )],
        experience=[ScoredExperienceEntry(
            title="Data Engineering Intern",
            company="Instacodigo",
            location="Remote",
            date_range=DateRange(start="Jul 2025", end="Oct 2025"),
            bullets=["Built ETL pipeline with Airflow"],
            relevance_score=10,
            relevance_reason="Direct match",
        )],
        projects=[ScoredProjectEntry(
            name="Hospital Information Website",
            description="Backend development course",
            tech_stack=["Node.js", "Express"],
            bullets=["Built REST API with JWT auth"],
            relevance_score=9,
            relevance_reason="Backend project",
            date_range=DateRange(start="2025"),
        )],
    )


@pytest.fixture
def generator() -> DocxGenerator:
    return DocxGenerator()


class TestDocxGenerator:
    def test_returns_bytes(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_output_is_valid_docx(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        # if this doesn't raise, it's a valid docx
        doc = Document(BytesIO(result))
        assert doc is not None

    def test_full_name_in_output(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        doc = Document(BytesIO(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Moaaz Emam Ahmed" in full_text

    def test_experience_entry_in_output(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        doc = Document(BytesIO(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Data Engineering Intern" in full_text
        assert "Instacodigo" in full_text

    def test_project_entry_in_output(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        doc = Document(BytesIO(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Hospital Information Website" in full_text

    def test_skills_in_output(self, generator, sample_tailored_cv):
        result = generator.generate(sample_tailored_cv)
        doc = Document(BytesIO(result))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Programming Languages" in full_text
        assert "Python" in full_text

    def test_content_type(self, generator):
        assert "wordprocessingml" in generator.content_type()

    def test_file_extension(self, generator):
        assert generator.file_extension() == "docx"

    def test_filename_format(self, generator, sample_tailored_cv):
        name = generator.filename(sample_tailored_cv)
        assert name == "moaaz_emam_ahmed_acme_corp.docx"

    def test_no_experience_still_generates(self, generator, sample_tailored_cv):
        sample_tailored_cv.experience = []
        result = generator.generate(sample_tailored_cv)
        assert isinstance(result, bytes)
        assert len(result) > 0